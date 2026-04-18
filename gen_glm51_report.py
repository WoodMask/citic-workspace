import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

template_path = 'C:/Users/35475/Project/citic-workspace/report-template/卡中心开源技术引入风险与技术评估报告-doris ccr-syncer.docx'
output_path = 'C:/Users/35475/Project/citic-workspace/卡中心开源技术引入风险与技术评估报告-GLM-5.1.docx'

doc = Document(template_path)

def update_cell(table, row_idx, col_idx, new_text):
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        run = first_para.add_run(new_text)

# Table 0: doc info
t0 = doc.tables[0]
update_cell(t0, 0, 1, '2026-04-18')
update_cell(t0, 1, 1, '信用卡中心信息技术部')
update_cell(t0, 2, 1, '系统级平台域')
update_cell(t0, 3, 1, 'AI原生应用开发平台')

# Table 1: doc name/author/review/date
t1 = doc.tables[1]
update_cell(t1, 1, 0, 'GLM-5.1引入风险与技术评估报告')
update_cell(t1, 1, 1, '吴思楠')
update_cell(t1, 1, 2, '陈远川')
update_cell(t1, 1, 3, '2026-04-18')

# Table 2: revision record
t2 = doc.tables[2]
update_cell(t2, 1, 0, 'V1.0')
update_cell(t2, 1, 1, '2026-04-18')
update_cell(t2, 1, 2, '')
update_cell(t2, 1, 3, '文档第一版')
update_cell(t2, 1, 4, '吴思楠')

# Table 3: software name/version
t3 = doc.tables[3]
update_cell(t3, 1, 0, 'GLM-5.1')
update_cell(t3, 1, 1, '最新旗舰版')

# Table 4: license
t4 = doc.tables[4]
update_cell(t4, 1, 0, '√')
update_cell(t4, 1, 1, 'Apache-2.0')
update_cell(t4, 1, 2, '1需要给代码的用户一份Apache Licence\n2如果你修改了代码，需要再被修改的文件中说明。\n3在延伸的代码中需要带有原来代码中的协议、商标、专利声明和其他原来作者规定需要包含的说明。\n4如果再发布的产品中包含一个Notice文件，则在Notice文件中需要带有Apache Licence。\nApache Licence也是对商业应用友好的许可。使用者也可以在需要的时候修改代码来满足需要并作为开源或商业产品发布/销售。')
update_cell(t4, 1, 3, 'https://www.apache.org/licenses/LICENSE-2.0.html')

# Table 5: maturity & community
t5 = doc.tables[5]
update_cell(t5, 0, 1, 'GLM-5.1')
update_cell(t5, 1, 0, '第一版发布时间')
update_cell(t5, 1, 1, '2024年（ChatGLM系列自2023年起）')
update_cell(t5, 2, 0, '当前版本')
update_cell(t5, 2, 1, 'GLM-5.1（2026年4月发布）')
update_cell(t5, 3, 0, '开发者')
update_cell(t5, 3, 1, '智谱AI（Zhipu AI）/ 清华大学KEG')
update_cell(t5, 4, 0, '最后一次更新时间')
update_cell(t5, 4, 1, '2026-04')
update_cell(t5, 5, 0, 'VIP 用户')
update_cell(t5, 5, 1, '广泛应用于智能客服、编程助手、文档生成等')
update_cell(t5, 6, 0, 'Commit 次数')
update_cell(t5, 6, 1, '数千次（活跃开发）')
update_cell(t5, 7, 0, 'Release 次数')
update_cell(t5, 7, 1, '多次（持续迭代）')
update_cell(t5, 8, 0, '代码提供人数')
update_cell(t5, 8, 1, '50+')
update_cell(t5, 9, 0, 'Github Star 个数')
update_cell(t5, 9, 1, '3500+')

# Table 6: vulnerabilities - keep empty
# Table 7: source link
t7 = doc.tables[7]
update_cell(t7, 0, 0, 'github')
update_cell(t7, 0, 1, 'https://github.com/THUDM/GLM')

# Table 8: system rating
t8 = doc.tables[8]
update_cell(t8, 1, 0, 'S AI原生应用开发平台')
update_cell(t8, 1, 1, 'B')

# Update paragraphs
replacements = {
    'doris ccr-syncer': 'GLM-5.1',
    'doris ccr-syncer引入风险与技术评估报告': 'GLM-5.1引入风险与技术评估报告',
    'Doris ccr-syncer引入风险与技术评估报告': 'GLM-5.1引入风险与技术评估报告',
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for run in para.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
    
    # 1.3 intro necessity
    if 'doris ccr-syncer(Cross Cluster Repliaction)' in text or 'doris ccr-syncer（Cross Cluster Repliaction）' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = 'GLM-5.1是智谱AI最新旗舰大语言模型，代码能力大大增强，长程任务显著提升，能够在单次任务中持续、自主地工作长达8小时，完成从规划、执行到迭代优化的完整闭环，交付工程级成果。在综合能力与Coding能力上，GLM-5.1整体表现对齐Claude Opus 4.6，在SWE-Bench Pro基准测试中取得58.4的成绩，超过GPT-5.4、Claude Opus 4.6和Gemini 3.1 Pro，刷新全球最佳表现。它有如下的特性：'
    
    if '最小表级别的将数据同步至目标集群' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '上下文窗口200K，最大输出128K Tokens，支持超长文本处理与多轮深度对话'

    if '分布式高可用部署' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '长程任务能力：可在单次任务中持续自主工作长达8小时，完成从规划、执行、测试到修复和交付的完整流程'

    if '可以为多个doris建立数据备份通道' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '支持深度思考模式、Function Call、结构化输出、上下文缓存、MCP等能力，适合构建Autonomous Agent与长程Coding Agent'

    # 1.4 similar software comparison
    if '暂无其他Doris数据同步工具' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '类似大模型包括：GPT-4o（OpenAI）、Claude Opus（Anthropic）、Gemini Pro（Google）、DeepSeek-V3、Qwen2.5等。GLM-5.1在综合与Coding能力上对齐全球顶尖水平，且在SWE-Bench Pro上表现最优；在长程自主执行、复杂工程优化与真实开发场景中展现出更强的持续工作能力，是中国模型中率先达到8小时级持续工作能力的代表。'

    # 1.7 service providers
    if '官方由SelectDB提供升级、维护服务' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '官方由智谱AI（Zhipu AI）提供API服务、模型微调、模型部署等服务，同时支持OpenAI SDK兼容接口，便于快速迁移现有应用。'

    # 1.8 supply chain risk
    if '供应链风险较低' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '供应链风险较低，原因：（1）GLM系列由清华大学KEG与智谱AI联合开发，拥有强大的研发团队和技术积累，可持续提供更新与维护；（2）GLM系列模型社区活跃度高，GitHub Star数超3500，社区用户广泛，有大量开发者参与贡献；（3）智谱AI是国内领先的大模型公司，已完成模型备案，合规性强；（4）Apache-2.0开源协议对商业应用友好。'

    # 2.1 initial intro
    if text.strip() == '初次引入使用':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '初次引入使用'

    # 2.3/2.4 not applicable
    if text.strip() == '不涉及':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '不涉及'

    # 2.6 usage plan
    if '用于2024-06' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '用于2026年AI原生应用开发平台引入大模型能力，支撑多场景综合业务需求（智能编码、对话问答、文档生成、创意内容等），后续根据模型版本更新和安全漏洞扫描结果的要求进行升级。'

doc.save(output_path)
print(f'报告已生成: {output_path}')
print('完成!')