import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

template_path = 'C:/Users/35475/Project/citic-workspace/report-template/卡中心开源技术引入风险与技术评估报告-doris ccr-syncer.docx'
output_path = 'C:/Users/35475/Project/citic-workspace/卡中心开源技术引入风险与技术评估报告-MiniMax-M2.7.docx'

doc = Document(template_path)

def update_cell(table, row_idx, col_idx, new_text):
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        run = first_para.add_run(new_text)

# Table 0: doc info
t0 = doc.tables[0]
update_cell(t0, 0, 1, '2026-04-18')
update_cell(t0, 1, 1, '信用卡中心信息技术部')
update_cell(t0, 2, 1, '系统级平台域')
update_cell(t0, 3, 1, 'AI原生应用开发平台')

# Table 1: doc name/author/review/date
t1 = doc.tables[1]
update_cell(t1, 1, 0, 'MiniMax-M2.7引入风险与技术评估报告')
update_cell(t1, 1, 1, '吴思楠')
update_cell(t1, 1, 2, '陈远川')
update_cell(t1, 1, 3, '2026-04-18')

# Table 2: revision record
t2 = doc.tables[2]
update_cell(t2, 1, 0, 'V1.0')
update_cell(t2, 1, 1, '2026-04-18')
update_cell(t2, 1, 2, '')
update_cell(t2, 1, 3, '文档第一版')
update_cell(t2, 1, 4, '吴思楠')

# Table 3: software name/version
t3 = doc.tables[3]
update_cell(t3, 1, 0, 'MiniMax-M2.7')
update_cell(t3, 1, 1, '最新旗舰版')

# Table 4: license
t4 = doc.tables[4]
update_cell(t4, 1, 0, '√')
update_cell(t4, 1, 1, 'MIT')
update_cell(t4, 1, 2, 'MIT协议是最宽松的开源协议之一：\n1使用者可以任意处理该软件，包括使用、复制、修改、合并、出版发行、散布、再授权及贩售软件副本。\n2被授权人可根据程序修改软件，但修改后的软件必须保留原作者的版权声明。\n3MIT协议对商业应用友好，使用者可在需要时修改代码以满足需求并作为开源或商业产品发布/销售。')
update_cell(t4, 1, 3, 'https://opensource.org/licenses/MIT')

# Table 5: maturity & community
t5 = doc.tables[5]
update_cell(t5, 0, 1, 'MiniMax-M2.7')
update_cell(t5, 1, 0, '第一版发布时间')
update_cell(t5, 1, 1, '2026年3月（MiniMax M2系列自2025年起）')
update_cell(t5, 2, 0, '当前版本')
update_cell(t5, 2, 1, 'MiniMax-M2.7（2026年3月发布）')
update_cell(t5, 3, 0, '开发者')
update_cell(t5, 3, 1, 'MiniMax AI（上海稀牛智能科技有限公司）')
update_cell(t5, 4, 0, '最后一次更新时间')
update_cell(t5, 4, 1, '2026-04')
update_cell(t5, 5, 0, 'VIP 用户')
update_cell(t5, 5, 1, '广泛应用于智能Agent、编程助手、办公自动化等')
update_cell(t5, 6, 0, 'Commit 次数')
update_cell(t5, 6, 1, '数十次（活跃开发）')
update_cell(t5, 7, 0, 'Release 次数')
update_cell(t5, 7, 1, '持续迭代')
update_cell(t5, 8, 0, '代码提供人数')
update_cell(t5, 8, 1, 'MiniMax研发团队')
update_cell(t5, 9, 0, 'Github Star 个数')
update_cell(t5, 9, 1, '248+')

# Table 7: source link
t7 = doc.tables[7]
update_cell(t7, 0, 0, 'github')
update_cell(t7, 0, 1, 'https://github.com/MiniMax-AI/MiniMax-M2.7')

# Table 8: system rating
t8 = doc.tables[8]
update_cell(t8, 1, 0, 'S AI原生应用开发平台')
update_cell(t8, 1, 1, 'B')

# Update paragraphs
replacements = {
    'doris ccr-syncer': 'MiniMax-M2.7',
    'doris ccr-syncer引入风险与技术评估报告': 'MiniMax-M2.7引入风险与技术评估报告',
    'Doris ccr-syncer引入风险与技术评估报告': 'MiniMax-M2.7引入风险与技术评估报告',
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for run in para.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
    
    # 1.3 intro necessity
    if 'doris ccr-syncer(Cross Cluster Repliaction)' in text or 'doris ccr-syncer（Cross Cluster Repliaction）' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = 'MiniMax-M2.7是MiniMax最新旗舰大语言模型，具备模型自我进化能力，可自行构建复杂Agent Harness，基于Agent Teams、复杂Skills、Tool Search等能力完成高度复杂的生产力任务。在软件工程领域表现优异，SWE-Pro基准测试得分56.22%，接近Opus水平；端到端项目交付能力VIBE-Pro得分55.6%。它有如下的特性：'
    
    if '最小表级别的将数据同步至目标集群' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '自我进化能力：模型可参与自身迭代优化，在MLE Bench Lite测试中取得66.6%得牌率，仅次于Opus-4.6和GPT-5.4'
    
    if '分布式高可用部署' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '强工程与Coding能力：在真实软件工程场景中表现优异，涵盖日志分析、Bug定位、代码重构、代码安全、机器学习等方向'
    
    if '可以为多个doris建立数据备份通道' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '专业办公能力：GDPval-AA评测ELO得分1495（开源最高），对Excel/PPT/Word复杂编辑能力显著提升，支持多轮高保真编辑'
    
    # 1.4 similar software comparison
    if '暂无其他Doris数据同步工具' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '类似大模型包括：GPT-4o（OpenAI）、Claude Opus/Sonnet（Anthropic）、Gemini Pro（Google）、DeepSeek-V3、GLM-5.1等。MiniMax-M2.7在软件工程基准测试上表现接近国际一线水平，SWE-Pro 56.22%接近GPT-5.3-Codex；在Agent能力方面，支持原生Agent Teams多智能体协作；在办公场景，GDPval-AA ELO 1495为开源最高。是国内模型中率先具备完整自我进化能力的代表。'
    
    # 1.7 service providers
    if '官方由SelectDB提供升级、维护服务' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '官方由MiniMax AI提供API服务、模型部署、Token Plan订阅等服务，同时支持Claude Code、Cursor、Cline等主流AI编程工具接入，便于快速迁移现有开发流程。'
    
    # 1.8 supply chain risk
    if '供应链风险较低' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '供应链风险较低，原因：（1）MiniMax是国内领先的通用人工智能科技公司，成立于2022年初，已完成模型备案，合规性强；（2）MiniMax自主研发多模态大模型，拥有强大的研发团队和技术积累；（3）MiniMax-M2系列模型在GitHub和社区获得积极反馈，服务全球超9亿个人用户、214000+企业客户；（4）MIT开源协议对商业应用友好。'
    
    # 2.1 initial intro
    if text.strip() == '初次引入使用':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '初次引入使用'
    
    # 2.3/2.4 not applicable
    if text.strip() == '不涉及':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '不涉及'
    
    # 2.6 usage plan
    if '用于2024-06' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '用于2026年AI原生应用开发平台引入大模型能力，支撑多场景综合业务需求（智能编码、Agent协作、办公自动化、日志分析等），后续根据模型版本更新和安全漏洞扫描结果的要求进行升级。'

doc.save(output_path)
print(f'报告已生成: {output_path}')
print('完成!')