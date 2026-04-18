"""
开源技术引入风险与技术评估报告生成脚本
根据模板和数据生成完整的评估报告文档
"""
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Optional


def update_cell(table, row_idx: int, col_idx: int, new_text: str):
    """
    更新表格单元格内容，保留原有格式
    
    Args:
        table: docx表格对象
        row_idx: 行索引
        col_idx: 列索引
        new_text: 新文本内容
    """
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        run = first_para.add_run(new_text)


def update_tables(doc: Document, data: Dict):
    """
    更新文档中的所有表格
    
    Args:
        doc: Document对象
        data: 报告数据字典
    """
    tables = doc.tables
    
    if len(tables) < 9:
        raise ValueError("模板表格数量不足，需要9个表格")
    
    # Table 0: 文档基本信息
    t0 = tables[0]
    update_cell(t0, 0, 1, data.get('release_date', ''))
    update_cell(t0, 1, 1, data.get('department', ''))
    update_cell(t0, 2, 1, data.get('domain', ''))
    update_cell(t0, 3, 1, data.get('project', ''))
    
    # Table 1: 文档名称
    t1 = tables[1]
    report_title = f"{data.get('software_name', '')}引入风险与技术评估报告"
    update_cell(t1, 1, 0, report_title)
    update_cell(t1, 1, 1, data.get('author', ''))
    update_cell(t1, 1, 2, data.get('reviewer', ''))
    update_cell(t1, 1, 3, data.get('release_date', ''))
    
    # Table 2: 修订记录
    t2 = tables[2]
    update_cell(t2, 1, 0, 'V1.0')
    update_cell(t2, 1, 1, data.get('release_date', ''))
    update_cell(t2, 1, 2, '')
    update_cell(t2, 1, 3, '文档第一版')
    update_cell(t2, 1, 4, data.get('author', ''))
    
    # Table 3: 软件名称和版本
    t3 = tables[3]
    update_cell(t3, 1, 0, data.get('software_name', ''))
    update_cell(t3, 1, 1, data.get('version', ''))
    
    # Table 4: 开源协议
    t4 = tables[4]
    update_cell(t4, 1, 0, '√')
    update_cell(t4, 1, 1, data.get('license', ''))
    license_notes = get_license_notes(data.get('license', ''))
    update_cell(t4, 1, 2, license_notes)
    update_cell(t4, 1, 3, get_license_url(data.get('license', '')))
    
    # Table 5: 成熟度和社区活跃度
    t5 = tables[5]
    update_cell(t5, 0, 1, data.get('software_name', ''))
    update_cell(t5, 1, 0, '第一版发布时间')
    update_cell(t5, 1, 1, data.get('first_release', ''))
    update_cell(t5, 2, 0, '当前版本')
    update_cell(t5, 2, 1, data.get('version', ''))
    update_cell(t5, 3, 0, '开发者')
    update_cell(t5, 3, 1, data.get('developer', ''))
    update_cell(t5, 4, 0, '最后一次更新时间')
    update_cell(t5, 4, 1, data.get('last_update', ''))
    update_cell(t5, 5, 0, 'VIP用户')
    update_cell(t5, 5, 1, data.get('vip_users', '广泛应用于相关场景'))
    update_cell(t5, 6, 0, 'Commit次数')
    update_cell(t5, 6, 1, data.get('commits', ''))
    update_cell(t5, 7, 0, 'Release次数')
    update_cell(t5, 7, 1, data.get('releases', '多次'))
    update_cell(t5, 8, 0, '代码提供人数')
    update_cell(t5, 8, 1, data.get('contributors', ''))
    update_cell(t5, 9, 0, 'Github Star个数')
    update_cell(t5, 9, 1, data.get('stars', ''))
    
    # Table 6: 安全漏洞 - 保持为空或填写已知漏洞
    t6 = tables[6]
    vulnerabilities = data.get('vulnerabilities', [])
    if vulnerabilities:
        for i, vuln in enumerate(vulnerabilities[:5]):  # 最多显示5个漏洞
            if i + 1 < len(t6.rows):
                update_cell(t6, i + 1, 0, vuln.get('id', ''))
                update_cell(t6, i + 1, 1, vuln.get('affected_version', ''))
                update_cell(t6, i + 1, 2, vuln.get('description', ''))
    
    # Table 7: 代码来源
    t7 = tables[7]
    update_cell(t7, 0, 0, 'github')
    update_cell(t7, 0, 1, data.get('github_url', ''))
    
    # Table 8: 系统评级
    t8 = tables[8]
    update_cell(t8, 1, 0, f"S {data.get('project', '')}")
    update_cell(t8, 1, 1, data.get('rating', ''))


def update_paragraphs(doc: Document, data: Dict):
    """
    更新文档中的段落内容
    
    Args:
        doc: Document对象
        data: 报告数据字典
    """
    software_name = data.get('software_name', '')
    
    replacements = {
        'GLM-5.1': software_name,
        'GLM-5.1引入风险与技术评估报告': f'{software_name}引入风险与技术评估报告',
        '智谱AI': data.get('developer', ''),
        '智谱AI（Zhipu AI）': data.get('developer', ''),
        '清华大学KEG与智谱AI': data.get('developer', ''),
        '清华大学KEG': '',
        'Zhipu AI': data.get('developer', ''),
    }
    
    necessity_text = data.get('necessity', '')
    comparison_text = data.get('comparison', '')
    service_provider_text = data.get('service_provider', '')
    supply_chain_risk_text = data.get('supply_chain_risk', '')
    usage_plan_text = data.get('usage_plan', '')
    features_text = data.get('features', [])
    
    for para in doc.paragraphs:
        text = para.text
        
        # 简单文本替换
        for run in para.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
        
        # 1.3 引入必要性评估
        if 'GLM-5.1是智谱AI最新旗舰大语言模型' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = necessity_text
        
        # 特性列表项替换
        if '上下文窗口200K，最大输出128K' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = features_text[0] if features_text else ''
        
        if '长程任务能力：可在单次任务中持续自主工作' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = features_text[1] if len(features_text) > 1 else ''
        
        if '支持深度思考模式、Function Call' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = features_text[2] if len(features_text) > 2 else ''
        
        # 1.4 类似开源软件对比
        if '类似大模型包括' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = comparison_text
        
        # 1.7 服务供应商
        if '官方由智谱AI' in text or ('官方由' in text and '提供API服务' in text):
            for run in para.runs:
                run.text = ''
            para.runs[0].text = service_provider_text
        
        # 1.8 供应链风险评估
        if '供应链风险较低，原因' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = supply_chain_risk_text
        
        # 2.1 是否初次引入
        if text.strip() == '初次引入使用':
            for run in para.runs:
                run.text = ''
            para.runs[0].text = '初次引入使用' if data.get('is_initial', True) else '版本升级'
        
        # 2.3/2.4 不涉及
        if text.strip() == '不涉及':
            pass
        
        # 2.6 使用计划
        if '用于2026年AI原生应用开发平台引入大模型能力' in text:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = usage_plan_text


def get_license_notes(license_type: str) -> str:
    """
    获取开源协议注意事项说明
    
    Args:
        license_type: 协议类型
    
    Returns:
        协议注意事项文本
    """
    notes = {
        'Apache-2.0': '''1需要给代码的用户一份Apache Licence 
2如果你修改了代码，需要再被修改的文件中说明。
3在延伸的代码中需要带有原来代码中的协议、商标、专利声明和其他原来作者规定需要包含的说明。
4如果再发布的产品中包含一个Notice文件，则在Notice文件中需要带有Apache Licence。
Apache Licence也是对商业应用友好的许可。使用者也可以在需要的时候修改代码来满足需要并作为开源或商业产品发布/销售。''',
        
        'MIT': '''MIT许可证非常简洁宽松，仅要求保留版权声明和许可声明。
允许任意处理代码，包括使用、复制、修改、合并、出版发行、散布、再授权及贩售软件的副本。
对商业应用友好，无额外限制。''',
        
        'GPL': '''GPL许可证要求：
1任何使用、修改或衍生的代码都必须以GPL协议发布。
2必须保持开源，不能闭源商业使用。
3修改后的代码必须开源并保留原作者版权声明。
适合开源项目，不适合商业闭源产品。''',
        
        'LGPL': '''LGPL许可证允许：
1动态链接使用LGPL库的程序可以不开源。
2修改LGPL库本身的代码必须以LGPL协议开源。
3适合商业产品引用开源库的场景。'''
    }
    return notes.get(license_type, '')


def get_license_url(license_type: str) -> str:
    """
    获取开源协议官方链接
    
    Args:
        license_type: 协议类型
    
    Returns:
        协议官方链接
    """
    urls = {
        'Apache-2.0': 'https://www.apache.org/licenses/LICENSE-2.0.html',
        'MIT': 'https://opensource.org/licenses/MIT',
        'GPL': 'https://www.gnu.org/licenses/gpl-3.0.html',
        'LGPL': 'https://www.gnu.org/licenses/lgpl-3.0.html',
    }
    return urls.get(license_type, '')


def generate_report(
    template_path: str,
    output_path: str,
    data: Dict
) -> str:
    """
    基于模板生成开源技术评估报告
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        data: 报告数据字典，包含所有必要字段
    
    Returns:
        生成的报告文件路径
    
    Raises:
        FileNotFoundError: 模板文件不存在
        ValueError: 数据不完整或模板格式错误
    """
    # 加载模板
    doc = Document(template_path)
    
    # 更新表格
    update_tables(doc, data)
    
    # 更新段落
    update_paragraphs(doc, data)
    
    # 保存报告
    doc.save(output_path)
    
    return output_path


def create_report_data_template() -> Dict:
    """
    创建报告数据模板字典
    
    Returns:
        包含所有字段的空模板字典
    """
    return {
        # 基本信息
        'software_name': '',
        'version': '',
        'release_date': '',
        'department': '信用卡中心信息技术部',
        'domain': '',
        'project': '',
        'rating': '',
        'author': '',
        'reviewer': '',
        
        # 技术信息
        'license': '',
        'developer': '',
        'github_url': '',
        'first_release': '',
        'last_update': '',
        'stars': '',
        'commits': '',
        'contributors': '',
        'vip_users': '',
        'releases': '',
        
        # 业务信息
        'necessity': '',
        'features': [],
        'comparison': '',
        'service_provider': '',
        'supply_chain_risk': '',
        'is_initial': True,
        'usage_plan': '',
        
        # 可选信息
        'vulnerabilities': [],
    }


if __name__ == '__main__':
    # 示例调用
    skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    template = os.path.join(skill_dir, 'template', '卡中心开源技术引入风险与技术评估报告-GLM-5.1.docx')
    output = '卡中心开源技术引入风险与技术评估报告-示例.docx'
    
    data = create_report_data_template()
    data['software_name'] = '示例软件'
    data['version'] = 'v1.0'
    data['release_date'] = '2026-04-18'
    
    result = generate_report(template, output, data)
    print(f'报告已生成: {result}')