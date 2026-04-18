import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

template_path = 'C:/Users/35475/Project/citic-workspace/report-template/卡中心开源技术引入风险与技术评估报告-doris ccr-syncer.docx'
output_path = 'C:/Users/35475/Project/citic-workspace/卡中心开源技术引入风险与技术评估报告-Gemma-4-31B.docx'

doc = Document(template_path)

def update_cell(table, row_idx, col_idx, new_text):
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        run = first_para.add_run(new_text)

def clear_para_and_set(para, new_text):
    for run in para.runs:
        run.text = ''
    para.runs[0].text = new_text

# Table 0: doc info
t0 = doc.tables[0]
update_cell(t0, 0, 1, '2026-04-19')
update_cell(t0, 1, 1, '信用卡中心信息技术部')
update_cell(t0, 2, 1, '系统级平台域')
update_cell(t0, 3, 1, 'AI原生应用开发平台')

# Table 1: doc name/author/review/date
t1 = doc.tables[1]
update_cell(t1, 1, 0, 'Gemma-4-31B引入风险与技术评估报告')
update_cell(t1, 1, 1, '吴思楠')
update_cell(t1, 1, 2, '陈远川')
update_cell(t1, 1, 3, '2026-04-19')

# Table 2: revision record
t2 = doc.tables[2]
update_cell(t2, 1, 0, 'V1.0')
update_cell(t2, 1, 1, '2026-04-19')
update_cell(t2, 1, 2, '')
update_cell(t2, 1, 3, '文档第一版')
update_cell(t2, 1, 4, '吴思楠')

# Table 3: software name/version
t3 = doc.tables[3]
update_cell(t3, 1, 0, 'Gemma-4-31B')
update_cell(t3, 1, 1, '31B Dense 指令微调版')

# Table 4: license
t4 = doc.tables[4]
update_cell(t4, 1, 0, '√')
update_cell(t4, 1, 1, 'Apache-2.0')
update_cell(t4, 1, 2, '1需要给代码的用户一份Apache Licence\n2如果你修改了代码，需要再被修改的文件中说明。\n3在延伸的代码中需要带有原来代码中的协议、商标、专利声明和其他原来作者规定需要包含的说明。\n4如果再发布的产品中包含一个Notice文件，则在Notice文件中需要带有Apache Licence。\nApache Licence也是对商业应用友好的许可。使用者也可以在需要的时候修改代码来满足需要并作为开源或商业产品发布/销售。')
update_cell(t4, 1, 3, 'https://www.apache.org/licenses/LICENSE-2.0.html')

# Table 5: maturity & community
t5 = doc.tables[5]
update_cell(t5, 0, 1, 'Gemma-4-31B')
update_cell(t5, 1, 0, '第一版发布时间')
update_cell(t5, 1, 1, '2023年（Gemma系列），Gemma 4于2026年发布')
update_cell(t5, 2, 0, '当前版本')
update_cell(t5, 2, 1, 'Gemma-4-31B-it（2026年4月发布）')
update_cell(t5, 3, 0, '开发者')
update_cell(t5, 3, 1, 'Google DeepMind')
update_cell(t5, 4, 0, '最后一次更新时间')
update_cell(t5, 4, 1, '2026-04')
update_cell(t5, 5, 0, 'VIP 用户')
update_cell(t5, 5, 1, '广泛应用于编程助手、多模态理解、Agentic应用等')
update_cell(t5, 6, 0, 'Commit 次数')
update_cell(t5, 6, 1, '活跃开发（持续迭代）')
update_cell(t5, 7, 0, 'Release 次数')
update_cell(t5, 7, 1, '多次（持续迭代）')
update_cell(t5, 8, 0, '代码提供人数')
update_cell(t5, 8, 1, 'Google DeepMind团队+社区贡献者')
update_cell(t5, 9, 0, 'Github Star 个数')
update_cell(t5, 9, 1, 'ModelScope下载量9266+')

# Table 6: vulnerabilities - keep empty/no known vulnerabilities
t6 = doc.tables[6]

# Table 7: source link
t7 = doc.tables[7]
update_cell(t7, 0, 0, 'github')
update_cell(t7, 0, 1, 'https://github.com/google-gemma')
if len(t7.rows) > 1:
    update_cell(t7, 1, 0, 'modelscope')
    update_cell(t7, 1, 1, 'https://www.modelscope.cn/models/google/gemma-4-31B')

# Table 8: system rating
t8 = doc.tables[8]
update_cell(t8, 1, 0, 'S AI原生应用开发平台')
update_cell(t8, 1, 1, 'B')

# Update paragraphs
replacements = {
    'doris ccr-syncer': 'Gemma-4-31B',
    'doris ccr-syncer引入风险与技术评估报告': 'Gemma-4-31B引入风险与技术评估报告',
    'Doris ccr-syncer引入风险与技术评估报告': 'Gemma-4-31B引入风险与技术评估报告',
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for run in para.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
    
    if 'doris ccr-syncer(Cross Cluster Repliaction)' in text or 'doris ccr-syncer（Cross Cluster Repliaction）' in text:
        clear_para_and_set(para, 'Gemma-4-31B是Google DeepMind开发的开源多模态大语言模型，是Gemma系列的最新版本。Gemma 4模型处理文本和图像输入（小模型还支持音频），生成文本输出，支持高达256K tokens的上下文窗口，并保持140多种语言的多语言支持。Gemma 4采用Dense和MoE两种架构，在推理、编程和多模态理解方面表现出色。31B Dense版本在多项基准测试中表现优异：MMLU Pro 85.2%、LiveCodeBench v6 80.0%、AIME 2026 89.2%、GPQA Diamond 84.3%。它有如下的特性：')
    
    if '最小表级别的将数据同步至目标集群' in text:
        clear_para_and_set(para, '256K长上下文窗口：支持超长文本处理与多轮深度对话，适合处理复杂文档、代码库分析等长程任务')
    
    if '分布式高可用部署' in text:
        clear_para_and_set(para, 'Thinking推理模式：内置可配置的逐步推理能力，支持复杂问题的深度分析与逻辑推理')
    
    if '可以为多个doris建立数据备份通道' in text:
        clear_para_and_set(para, '原生Function Calling支持：支持结构化工具调用，适合构建Autonomous Agent与Agentic应用，实现复杂任务的自主规划与执行')
    
    if '暂无其他Doris数据同步工具' in text:
        clear_para_and_set(para, '类似大模型包括：开源模型GLM-5.1（智谱AI）、Qwen3.6-35B-A3B（阿里）、DeepSeek-V3（深度求索）；闭源模型GPT-4o（OpenAI）、Claude Opus（Anthropic）、Gemini Pro（Google）；同系列Gemma 3 27B。Gemma-4-31B相较于Gemma 3 27B在多项基准测试上显著提升：MMLU Pro从67.6%提升至85.2%（+17.6%），LiveCodeBench从29.1%提升至80.0%（+50.9%），AIME从20.8%提升至89.2%（+68.4%）。在开源模型中，Gemma-4-31B的编程能力与推理能力处于领先水平，且Apache-2.0协议对商业应用友好。')
    
    if '官方由SelectDB提供升级、维护服务' in text:
        clear_para_and_set(para, '官方由Google DeepMind提供模型权重、API服务及部署支持，可通过Hugging Face、ModelScope下载模型权重，支持Transformers、vLLM、llama.cpp等主流推理框架部署。第三方服务商包括各类云平台提供的模型托管与推理服务。')
    
    if '供应链风险较低' in text:
        clear_para_and_set(para, '供应链风险较低，原因：（1）Google DeepMind是全球顶尖AI研发团队，拥有强大的技术积累和持续研发能力；（2）Gemma系列模型社区活跃度高，ModelScope下载量超过9000次，GitHub社区用户广泛，有大量开发者参与贡献；（3）Apache-2.0开源协议对商业应用友好，无强制开源要求；（4）Google已完成模型备案，合规性强，符合AI安全标准。')
    
    if text.strip() == '初次引入使用':
        clear_para_and_set(para, '初次引入使用')
    
    if text.strip() == '不涉及':
        clear_para_and_set(para, '不涉及')
    
    if '用于2024-06' in text:
        clear_para_and_set(para, '用于2026年AI原生应用开发平台引入多模态大模型能力，支撑Agentic应用、编程助手、多模态理解等业务需求，后续根据模型版本更新和安全漏洞扫描结果的要求进行升级。')

doc.save(output_path)
print(f'报告已生成: {output_path}')
print('完成!')