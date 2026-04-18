import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

template_path = 'C:/Users/35475/Project/citic-workspace/report-template/卡中心开源技术引入风险与技术评估报告-doris ccr-syncer.docx'
output_path = 'C:/Users/35475/Project/citic-workspace/卡中心开源技术引入风险与技术评估报告-Qwen3.6-35B-A3B.docx'

doc = Document(template_path)

def update_cell(table, row_idx, col_idx, new_text):
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = new_text
    else:
        run = first_para.add_run(new_text)

t0 = doc.tables[0]
update_cell(t0, 0, 1, '2026-04-19')
update_cell(t0, 1, 1, '信用卡中心信息技术部')
update_cell(t0, 2, 1, '系统级平台域')
update_cell(t0, 3, 1, 'AI原生应用开发平台')

t1 = doc.tables[1]
update_cell(t1, 1, 0, 'Qwen3.6-35B-A3B引入风险与技术评估报告')
update_cell(t1, 1, 1, '吴思楠')
update_cell(t1, 1, 2, '陈远川')
update_cell(t1, 1, 3, '2026-04-19')

t2 = doc.tables[2]
update_cell(t2, 1, 0, 'V1.0')
update_cell(t2, 1, 1, '2026-04-19')
update_cell(t2, 1, 2, '')
update_cell(t2, 1, 3, '文档第一版')
update_cell(t2, 1, 4, '吴思楠')

t3 = doc.tables[3]
update_cell(t3, 1, 0, 'Qwen3.6-35B-A3B')
update_cell(t3, 1, 1, '最新版（2026年4月发布）')

t4 = doc.tables[4]
update_cell(t4, 1, 0, '√')
update_cell(t4, 1, 1, 'Apache-2.0')
update_cell(t4, 1, 2, '1需要给代码的用户一份Apache Licence\n2如果你修改了代码，需要再被修改的文件中说明。\n3在延伸的代码中需要带有原来代码中的协议、商标、专利声明和其他原来作者规定需要包含的说明。\n4如果再发布的产品中包含一个Notice文件，则在Notice文件中需要带有Apache Licence。\nApache Licence也是对商业应用友好的许可。使用者也可以在需要的时候修改代码来满足需要并作为开源或商业产品发布/销售。')
update_cell(t4, 1, 3, 'https://www.apache.org/licenses/LICENSE-2.0.html')

t5 = doc.tables[5]
update_cell(t5, 0, 1, 'Qwen3.6-35B-A3B')
update_cell(t5, 1, 0, '第一版发布时间')
update_cell(t5, 1, 1, '2026年4月（Qwen系列自2023年起）')
update_cell(t5, 2, 0, '当前版本')
update_cell(t5, 2, 1, 'Qwen3.6-35B-A3B（2026年4月发布）')
update_cell(t5, 3, 0, '开发者')
update_cell(t5, 3, 1, 'Qwen Team / Alibaba Group（阿里云）')
update_cell(t5, 4, 0, '最后一次更新时间')
update_cell(t5, 4, 1, '2026-04-16')
update_cell(t5, 5, 0, 'VIP 用户')
update_cell(t5, 5, 1, '广泛应用于智能编码、Agent开发、多语言应用等场景')
update_cell(t5, 6, 0, 'Commit 次数')
update_cell(t5, 6, 1, '7次（新仓库，活跃开发）')
update_cell(t5, 7, 0, 'Release 次数')
update_cell(t5, 7, 1, '持续迭代中')
update_cell(t5, 8, 0, '代码提供人数')
update_cell(t5, 8, 1, 'Qwen Team')
update_cell(t5, 9, 0, 'Github Star 个数')
update_cell(t5, 9, 1, '2800+')

t7 = doc.tables[7]
update_cell(t7, 0, 0, 'github')
update_cell(t7, 0, 1, 'https://github.com/QwenLM/Qwen3.6')

t8 = doc.tables[8]
update_cell(t8, 1, 0, 'S AI原生应用开发平台')
update_cell(t8, 1, 1, 'B')

replacements = {
    'doris ccr-syncer': 'Qwen3.6-35B-A3B',
    'doris ccr-syncer引入风险与技术评估报告': 'Qwen3.6-35B-A3B引入风险与技术评估报告',
    'Doris ccr-syncer引入风险与技术评估报告': 'Qwen3.6-35B-A3B引入风险与技术评估报告',
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for run in para.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
    
    if 'doris ccr-syncer(Cross Cluster Repliaction)' in text or 'doris ccr-syncer（Cross Cluster Repliation）' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = 'Qwen3.6-35B-A3B是阿里云Qwen团队发布的最新MoE（Mixture-of-Experts）大语言模型，总参数35B（350亿），激活参数3B（30亿），采用稀疏激活架构实现高效推理。该模型专注于Agentic Coding能力，在前端工作流和仓库级推理方面表现优异，同时支持Thinking Preservation功能，可保留跨对话历史的思考上下文。它有如下的特性：'
    
    if '最小表级别的将数据同步至目标集群' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = 'MoE架构高效推理：35B总参数/3B激活参数，推理效率媲美小模型，成本低延迟小'
    
    if '分布式高可用部署' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '超长上下文支持：256K原生上下文长度，可扩展至1M Tokens，适合处理大型代码仓库和长文档'
    
    if '可以为多个doris建立数据备份通道' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '多语言覆盖：支持201种语言和方言，全球部署能力强，适合国际化应用场景'
    
    if '暂无其他Doris数据同步工具' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '类似MoE大模型包括：Qwen3-30B-A3B（30B/3B）、Qwen3-235B-A22B（235B/22B）、DeepSeek-V3（671B/37B）、Mixtral-8x7B等。Qwen3.6-35B-A3B在Agentic Coding和Thinking Preservation方面有独特优势，专注于提升编码场景的实际效率；相比Qwen3-30B-A3B，参数规模更大、推理能力更强；Apache-2.0协议对商业应用友好，可自由定制和部署。'
    
    if '官方由SelectDB提供升级、维护服务' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '官方由阿里云Model Studio提供API服务、模型微调、模型部署等服务，同时支持OpenAI SDK兼容接口，便于快速迁移现有应用。可通过DashScope API或本地部署（SGLang、vLLM、llama.cpp等）使用。'
    
    if '供应链风险较低' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '供应链风险较低，原因：（1）Qwen系列由阿里云Qwen团队开发，拥有强大的研发团队和技术积累，可持续提供更新与维护；（2）Qwen系列模型社区活跃度高，GitHub Star数超2800，社区用户广泛，有大量开发者参与贡献；（3）阿里云是国内领先的云计算和大模型公司，已完成多项模型备案，合规性强；（4）Apache-2.0开源协议对商业应用友好，可自由使用、修改和分发。'
    
    if text.strip() == '初次引入使用':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '初次引入使用'
    
    if text.strip() == '不涉及':
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '不涉及'
    
    if '用于2024-06' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = '用于2026年AI原生应用开发平台引入大模型能力，重点支撑Agentic Coding、智能编码助手、多语言应用等场景，利用MoE架构实现高效推理和低成本部署，后续根据模型版本更新和安全漏洞扫描结果的要求进行升级。'

doc.save(output_path)
print(f'报告已生成: {output_path}')
print('完成!')